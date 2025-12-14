#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Flatten a mixed-format knowledge base into Markdown (or text/JSON) for model training.

Modes
- Folder mode: scan a whole folder (--root) and produce a consolidated Markdown file.
- Single-file mode: process one document (--in) and emit a single section to stdout by default
  (ideal for n8n / pipelines).

Key features
- Formats: TXT, MD, JSON, CSV, XLSX/XLSM, DOCX, PDF.
- Skips obvious binaries with a clear stub; verbose logging goes to stderr.
- Stable Table of Contents with robust, GitHub-like anchors (folder mode).
- Per-file sections separated by '---' and fenced blocks for safe parsing.
- Optional text normalisation (--clean-text) to standardise punctuation and Unicode.
- Optional citation stripping (--strip-citations) to remove tokens like '62\\u2020file:L1-L20'.
- Optional PII removal (--remove-pii) with basic or extended patterns.
- Optional per-section length cap (--max-section-chars) to avoid dataset imbalance.
- Optional JSONL sidecar for downstream pipelines.

Dependencies (optional by file type): pypdf, python-docx, pandas, openpyxl, tabulate

Notes on PII removal
- "basic" aims to be conservative: emails, phone numbers, and common labelled fields (e.g. "Name:").
- "extended" also targets IPs, MACs, UK NI numbers, IBANs, and payment cards (Luhn-checked).
- Regex-based redaction is imperfect. Treat it as a safety net, not a compliance guarantee.
"""

import argparse
import csv
import json
import sys
import re
import unicodedata
import hashlib
import os
from pathlib import Path
from datetime import datetime
from typing import List, Tuple, Optional, Dict, Any, Callable

# ---------------------------- optional imports ----------------------------
MISSING_DEPS: List[str] = []

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None
    MISSING_DEPS.append("pypdf")

try:
    import docx  # python-docx
except Exception:
    docx = None
    MISSING_DEPS.append("python-docx")

try:
    import pandas as pd
except Exception:
    pd = None
    MISSING_DEPS.append("pandas (and openpyxl)")

# ---------------------------- configuration ----------------------------
TEXT_EXTS = {".txt", ".md", ".markdown"}
JSON_EXTS = {".json"}
CSV_EXTS = {".csv"}
XLSX_EXTS = {".xlsx", ".xlsm"}
DOCX_EXTS = {".docx"}
PDF_EXTS = {".pdf"}
BINARY_HINT_EXTS = {
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff",
    ".mp4", ".mov", ".avi", ".mkv", ".mp3", ".wav",
    ".pptx", ".ppt", ".doc", ".xls", ".zip", ".rar", ".7z"
}
IGNORES = {".ds_store", "thumbs.db", "~$"}  # substring check for '~$' temp files

# ---------------------------- helpers ----------------------------

def eprint(msg: str) -> None:
    print(msg, file=sys.stderr)


def sha1_of_string(s: str) -> str:
    h = hashlib.sha1()
    h.update(s.encode("utf-8", errors="ignore"))
    return h.hexdigest()[:8]


def sha1_of_path(path: Path) -> str:
    # Stable id across runs, based on a normalised posix path string.
    return sha1_of_string(path.as_posix())


def norm_heading(text: str) -> str:
    text = text.replace("#", "").strip()
    return re.sub(r"\s+", " ", text)


def md_anchor(text: str) -> str:
    """GitHub-like anchor: lower, replace non-alnum with '-', collapse runs, strip edges."""
    t = text.strip().lower()
    t = re.sub(r"[^a-z0-9]+", "-", t)
    t = re.sub(r"-+", "-", t).strip("-")
    return t or "section"

# Unicode and punctuation normalisation
SMARTS = {
    "\u2018": "'", "\u2019": "'", "\u201C": '"', "\u201D": '"',
    "\u2013": "-",  # en dash
    "\u2014": "-",  # em dash
    "\u00A0": " ",  # non-breaking space
    "\u00AD": "",   # soft hyphen
    "\u2028": "\n", # line separator
    "\u2029": "\n", # paragraph separator
}


def clean_text(s: str) -> str:
    s = unicodedata.normalize("NFKC", s)
    for k, v in SMARTS.items():
        s = s.replace(k, v)
    # Remove other control characters except tab and newline.
    s = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]", "", s)
    return s


def is_ignored(name: str) -> bool:
    low = name.lower()
    return any(token in low for token in IGNORES)


def classify_ext(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in TEXT_EXTS:
        return "text"
    if ext in JSON_EXTS:
        return "json"
    if ext in CSV_EXTS:
        return "csv"
    if ext in XLSX_EXTS:
        return "xlsx"
    if ext in DOCX_EXTS:
        return "docx"
    if ext in PDF_EXTS:
        return "pdf"
    if ext in BINARY_HINT_EXTS:
        return "binary"
    return "unknown"


def strip_md_fences(s: str) -> str:
    """
    Remove a single outer fenced code block if present.
    This is used for JSONL/plain-text exports.
    """
    return re.sub(r"^```[a-zA-Z]*\n|\n```$", "", s, flags=re.DOTALL)


# ---------------------------- readers ----------------------------

def read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return path.read_text(encoding="latin-1", errors="replace")


def read_json_as_md(path: Path) -> str:
    try:
        obj = json.loads(read_text_file(path))
        # Keep JSON readable; exact keys/values may matter for training.
        pretty = json.dumps(obj, indent=2, ensure_ascii=False)
        return "```json\n" + pretty + "\n```"
    except Exception as e:
        return "```text\nFailed to parse JSON: %s\nRaw:\n%s\n```" % (e, read_text_file(path))


def read_csv_as_md(path: Path, max_rows: int = 5000) -> str:
    try:
        with path.open("r", encoding="utf-8", errors="replace", newline="") as f:
            reader = csv.reader(f)
            rows: List[List[str]] = []
            for i, row in enumerate(reader):
                if i >= max_rows:
                    break
                rows.append(row)
        if not rows:
            return "_Empty CSV_"
        header = rows[0]
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * len(header)) + " |",
        ]
        for r in rows[1:]:
            lines.append("| " + " | ".join(r) + " |")
        if len(rows) >= max_rows:
            lines.append("\n_Trimmed to first %d rows_" % max_rows)
        return "\n".join(lines)
    except Exception as e:
        return "```text\nFailed to read CSV: %s\n```" % e


def read_xlsx_as_md(path: Path, max_rows: int = 2000) -> str:
    if pd is None:
        return "```text\npandas/openpyxl not available.\n```"
    try:
        xl = pd.ExcelFile(path)
        parts: List[str] = []
        for sheet in xl.sheet_names:
            df = xl.parse(sheet, dtype=str).fillna("")
            tail = ""
            if len(df) > max_rows:
                df = df.head(max_rows)
                tail = "\n_Trimmed to first %d rows_" % max_rows
            md = df.to_markdown(index=False)
            parts.append("#### Sheet: %s\n\n%s%s" % (sheet, md, tail))
        return "\n\n".join(parts)
    except Exception as e:
        return "```text\nFailed to read XLSX: %s\n```" % e


def read_docx_as_text(path: Path) -> str:
    if docx is None:
        return "```text\npython-docx not available.\n```"
    try:
        d = docx.Document(str(path))
        paras = [p.text for p in d.paragraphs]
        text = "\n".join(paras).strip()
        if not text:
            tparts: List[str] = []
            for tbl in d.tables:
                for row in tbl.rows:
                    cells = [c.text for c in row.cells]
                    tparts.append("| " + " | ".join(cells) + " |")
            text = "\n".join(tparts)
        return text if text else "_Empty DOCX_"
    except Exception as e:
        return "```text\nFailed to read DOCX: %s\n```" % e


def read_pdf_as_text(path: Path, max_pages: int = 200) -> str:
    if PdfReader is None:
        return "```text\npypdf not available.\n```"
    try:
        reader = PdfReader(str(path))
        n = min(len(reader.pages), max_pages)
        pages: List[str] = []
        for i in range(n):
            try:
                pages.append(reader.pages[i].extract_text() or "")
            except Exception:
                pages.append("")
        body = "\n\n".join(pages).strip()
        if len(reader.pages) > max_pages:
            body += "\n\n_Trimmed to first %d pages_" % max_pages
        return body if body else "_Empty or image-only PDF_"
    except Exception as e:
        return "```text\nFailed to read PDF: %s\n```" % e


# ---------------------------- assembly ----------------------------

def build_toc(entries: List[Tuple[str, str]]) -> str:
    lines = ["## Table of Contents", ""]
    for title, anchor in entries:
        lines.append("- [%s](#%s)" % (title, anchor))
    lines.append("")
    return "\n".join(lines)


def extract_content(path: Path, kind: str, max_pdf_pages: int, clean: bool) -> str:
    if kind == "text":
        data = read_text_file(path)
        if clean:
            data = clean_text(data)
        return "```text\n" + data.rstrip() + "\n```"
    if kind == "json":
        return read_json_as_md(path)
    if kind == "csv":
        out = read_csv_as_md(path)
        return clean_text(out) if clean else out
    if kind == "xlsx":
        out = read_xlsx_as_md(path)
        return clean_text(out) if clean else out
    if kind == "docx":
        out = read_docx_as_text(path)
        if clean:
            out = clean_text(out)
        return "```text\n" + out.rstrip() + "\n```"
    if kind == "pdf":
        out = read_pdf_as_text(path, max_pages=max_pdf_pages)
        if clean:
            out = clean_text(out)
        return "```text\n" + out.rstrip() + "\n```"
    if kind == "binary":
        try:
            size = os.path.getsize(path)
        except Exception:
            size = -1
        return "_Binary/media file skipped (size=%d bytes)_" % size
    return "_Unknown file type skipped._"


# ---------------------------- PII removal ----------------------------

# Pre-compiled patterns. Keep strings ASCII; the regex engine understands \uXXXX escapes.
CIT_PAT = re.compile(r"\b\d+\u2020[^\s:]+:L\d+-L\d+\b")

EMAIL_PAT = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.I)

# Broad, but practical: optional country code, optional leading 0, then 9+ digits allowing spaces/hyphens.
PHONE_PAT = re.compile(r"\b(?:\+\d{1,3}\s?)?(?:\(?0\)?[\s-]?)?(?:\d[\s-]?){9,}\b")

IPV4_PAT = re.compile(r"\b(?:(?:25[0-5]|2[0-4]\d|1?\d?\d)\.){3}(?:25[0-5]|2[0-4]\d|1?\d?\d)\b")
IPV6_PAT = re.compile(r"\b(?:[0-9a-f]{0,4}:){2,7}[0-9a-f]{0,4}\b", re.I)
MAC_PAT = re.compile(r"\b(?:[0-9A-F]{2}[:-]){5}[0-9A-F]{2}\b", re.I)

# UK National Insurance Number (approximate, with common invalid prefixes excluded).
UK_NI_PAT = re.compile(
    r"\b(?!BG)(?!GB)(?!NK)(?!KN)(?!TN)(?!NT)(?!ZZ)[A-CEGHJ-PR-TW-Z]{2}\s?\d{2}\s?\d{2}\s?\d{2}\s?[A-D]\b",
    re.I,
)

IBAN_PAT = re.compile(r"\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b", re.I)

# Candidate payment card numbers (validated via Luhn in extended mode).
CARD_CANDIDATE_PAT = re.compile(r"\b(?:\d[ -]*?){13,19}\b")


LABEL_PLACEHOLDERS: Dict[str, str] = {
    "name": "[redacted-name]",
    "full name": "[redacted-name]",
    "first name": "[redacted-name]",
    "last name": "[redacted-name]",
    "email": "[redacted-email]",
    "e-mail": "[redacted-email]",
    "phone": "[redacted-phone]",
    "mobile": "[redacted-phone]",
    "telephone": "[redacted-phone]",
    "tel": "[redacted-phone]",
    "address": "[redacted-address]",
    "dob": "[redacted-dob]",
    "date of birth": "[redacted-dob]",
    "national insurance": "[redacted-ni]",
    "ni number": "[redacted-ni]",
    "passport": "[redacted-passport]",
    "passport number": "[redacted-passport]",
    "iban": "[redacted-iban]",
    "sort code": "[redacted-sortcode]",
    "account number": "[redacted-account]",
    "acct number": "[redacted-account]",
    "card": "[redacted-card]",
    "card number": "[redacted-card]",
    "credit card": "[redacted-card]",
    "credit card number": "[redacted-card]",
}

LABEL_LINE_PAT = re.compile(
    r"(?im)^(?P<prefix>\s*)(?P<label>"
    + r"|".join(re.escape(k) for k in sorted(LABEL_PLACEHOLDERS.keys(), key=len, reverse=True))
    + r")(?P<sep>\s*[:=]\s*)(?P<value>.+?)\s*$"
)


def luhn_is_valid(number: str) -> bool:
    digits = [int(c) for c in number if c.isdigit()]
    if len(digits) < 13 or len(digits) > 19:
        return False
    total = 0
    parity = len(digits) % 2
    for i, d in enumerate(digits):
        if i % 2 == parity:
            d = d * 2
            if d > 9:
                d -= 9
        total += d
    return (total % 10) == 0


def redact_labelled_fields(text: str) -> str:
    def repl(m: re.Match) -> str:
        label = m.group("label").lower()
        placeholder = LABEL_PLACEHOLDERS.get(label, "[redacted]")
        return m.group("prefix") + m.group("label") + m.group("sep") + placeholder
    return LABEL_LINE_PAT.sub(repl, text)


def redact_cards_luhn(text: str) -> str:
    def repl(m: re.Match) -> str:
        raw = m.group(0)
        digits = "".join(c for c in raw if c.isdigit())
        if luhn_is_valid(digits):
            return "[redacted-card]"
        return raw
    return CARD_CANDIDATE_PAT.sub(repl, text)


def remove_pii(text: str, level: str) -> str:
    # Always do labelled fields first: it is high signal and lowers false positives.
    out = redact_labelled_fields(text)

    # Email is usually unambiguous.
    out = EMAIL_PAT.sub("[redacted-email]", out)

    # Extended: run more specific patterns before phone, so card numbers do not get labelled as phones.
    if level == "extended":
        out = redact_cards_luhn(out)
        out = IPV4_PAT.sub("[redacted-ip]", out)
        out = IPV6_PAT.sub("[redacted-ip]", out)
        out = MAC_PAT.sub("[redacted-mac]", out)
        out = UK_NI_PAT.sub("[redacted-ni]", out)
        out = IBAN_PAT.sub("[redacted-iban]", out)

    # Phone last (broad pattern).
    out = PHONE_PAT.sub("[redacted-phone]", out)

    return out


# ---------------------------- section builders ----------------------------

def build_section(
    title: str,
    path_display: str,
    kind: str,
    size_b: int,
    mtime: str,
    text_block: str,
    include_divider: bool = True,
    include_metadata: bool = True,
) -> str:
    parts: List[str] = []
    if include_divider:
        parts.append("---\n")
    parts.append("## %s\n" % norm_heading(title))
    if include_metadata:
        parts.append("_Path_: `%s`  |  _Type_: `%s`  |  _Size_: `%s`  |  _Modified_: `%s`\n"
                     % (path_display, kind, size_b, mtime))
    parts.append(text_block.rstrip() + "\n")
    return "\n".join(parts)


def post_process_block(
    text_block: str,
    strip_citations: bool,
    remove_pii_flag: bool,
    pii_level: str,
    max_section_chars: int,
) -> str:
    out = text_block
    if strip_citations:
        out = CIT_PAT.sub("", out)
    if remove_pii_flag:
        out = remove_pii(out, pii_level)
    if max_section_chars and len(out) > max_section_chars:
        out = out[:max_section_chars] + "\n\n_Trimmed for training export_"
    return out


def build_json_record(
    path_display: str,
    kind: str,
    sha: str,
    mtime: str,
    size_b: int,
    markdown: str,
    plain: str,
) -> Dict[str, Any]:
    return {
        "path": path_display,
        "type": kind,
        "sha": sha,
        "modified": mtime,
        "size": size_b,
        "markdown": markdown,
        "text": plain,
    }


def write_output(out_path: Optional[str], content: str, append: bool) -> None:
    if out_path is None or out_path == "-" or out_path.strip() == "":
        sys.stdout.write(content)
        if not content.endswith("\n"):
            sys.stdout.write("\n")
        return

    p = Path(out_path).resolve()
    p.parent.mkdir(parents=True, exist_ok=True)
    mode = "a" if append else "w"
    with p.open(mode, encoding="utf-8") as f:
        f.write(content if content.endswith("\n") else content + "\n")


# ---------------------------- main ----------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Flatten knowledge base to Markdown for training.")

    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--root", help="Root folder to scan (folder mode)")
    mode.add_argument("--in", dest="in_path", help="Input file to process (single-file mode)")

    parser.add_argument("--out", default=None, help="Output path. Use '-' for stdout. Defaults: root mode -> thintech_training_data.md, single mode -> stdout.")
    parser.add_argument("--append", action="store_true", help="Append to --out instead of overwriting (useful for single-file mode).")

    parser.add_argument("--emit", choices=["md", "text", "json"], default="md",
                        help="Single-file mode output format: md (section), text (plain), or json (record).")

    parser.add_argument("--max-pdf-pages", type=int, default=200, help="Safety limit for PDF page extraction")
    parser.add_argument("--include-unknown", action="store_true", help="Emit stubs for unknown extensions")
    parser.add_argument("--verbose", action="store_true", help="Enable detailed logging (stderr)")
    parser.add_argument("--clean-text", action="store_true", help="Normalise Unicode and typographic characters for training")
    parser.add_argument("--strip-citations", action="store_true", help="Remove legacy inline citation artefacts like '62\\u2020file:L1-L20'")
    parser.add_argument("--remove-pii", "--redact", dest="remove_pii", action="store_true",
                        help="Remove PII from output (emails/phones at minimum)")
    parser.add_argument("--pii-level", choices=["basic", "extended"], default="basic",
                        help="PII removal level. 'basic' is conservative; 'extended' is broader.")
    parser.add_argument("--max-section-chars", type=int, default=0, help="If > 0, trim each section's text to this many characters")
    parser.add_argument("--jsonl-out", type=str, default=None, help="Optional JSONL file to write one record per file (or per run in single-file mode).")

    args = parser.parse_args()

    # Decide output default by mode.
    out_path = args.out
    if out_path is None:
        out_path = "thintech_training_data.md" if args.root else "-"

    # Folder mode
    if args.root:
        root = Path(args.root).resolve()
        if not root.exists() or not root.is_dir():
            eprint("[ERROR] Root folder not found: %s" % root)
            sys.exit(1)

        if args.verbose:
            eprint("[INFO] Root:   %s" % root)
            eprint("[INFO] Output: %s" % Path(out_path).resolve() if out_path != "-" else "[stdout]")
            if MISSING_DEPS:
                eprint("[WARN] Missing optional dependencies:")
                for d in MISSING_DEPS:
                    eprint("       - %s" % d)

        # Collect files
        files: List[Path] = []
        for p in root.rglob("*"):
            if not p.is_file():
                continue
            if is_ignored(p.name):
                if args.verbose:
                    eprint("[SKIP ] Ignored: %s" % p.relative_to(root).as_posix())
                continue
            files.append(p)

        if not files:
            empty = "# Empty export\n\n_No files found in: %s_\n" % root
            write_output(out_path, empty, append=False)
            if args.verbose:
                eprint("[DONE] No files found under %s" % root)
            sys.exit(0)

        files.sort(key=lambda p: p.relative_to(root).as_posix().lower())
        if args.verbose:
            eprint("[INFO] Found %d files" % len(files))

        # Optional JSONL writer
        jsonl_fp: Optional[object] = None
        jsonl_path: Optional[Path] = None
        if args.jsonl_out:
            jsonl_path = Path(args.jsonl_out).resolve()
            jsonl_path.parent.mkdir(parents=True, exist_ok=True)
            jsonl_fp = jsonl_path.open("w", encoding="utf-8")
            if args.verbose:
                eprint("[INFO] JSONL:  %s" % jsonl_path)

        # Header and TOC
        lines: List[str] = []
        title = "Thintech Knowledge Base - Consolidated Training Document"
        ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%SZ")
        lines.append("# %s\n\n_Generated: %s_\n" % (title, ts))

        toc_entries: List[Tuple[str, str]] = []
        for p in files:
            rel = p.relative_to(root).as_posix()
            toc_entries.append((rel, md_anchor(rel)))
        lines.append(build_toc(toc_entries))

        processed = 0
        errors = 0

        for p in files:
            rel = p.relative_to(root).as_posix()
            kind = classify_ext(p)

            if kind == "unknown" and not args.include_unknown:
                if args.verbose:
                    eprint("[SKIP ] %s (unknown ext)" % rel)
                continue

            if args.verbose:
                eprint("[PROC ] %s (%s)" % (rel, kind))

            # Minimal metadata
            try:
                st = p.stat()
                mtime = datetime.utcfromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M:%SZ")
                size_b = st.st_size
            except Exception:
                mtime = ""
                size_b = -1

            # Extract
            try:
                content = extract_content(p, kind, args.max_pdf_pages, args.clean_text)
            except Exception as e:
                errors += 1
                content = "```text\nExtraction error: %s\n```" % e
                if args.verbose:
                    eprint("[ERROR] %s: %s" % (rel, e))

            text_block = post_process_block(
                content,
                strip_citations=args.strip_citations,
                remove_pii_flag=args.remove_pii,
                pii_level=args.pii_level,
                max_section_chars=args.max_section_chars,
            )

            section = build_section(
                title=rel,
                path_display=rel,
                kind=kind,
                size_b=size_b,
                mtime=mtime,
                text_block=text_block,
                include_divider=True,
                include_metadata=True,
            )
            lines.append(section)
            processed += 1

            if jsonl_fp is not None:
                plain = text_block
                if kind in {"text", "docx", "pdf", "json"}:
                    plain = strip_md_fences(plain)
                rec = {
                    "path": rel,
                    "type": kind,
                    "sha": sha1_of_path(Path(rel)),
                    "modified": mtime,
                    "size": size_b,
                    "text": plain,
                }
                jsonl_fp.write(json.dumps(rec, ensure_ascii=False) + "\n")

        if jsonl_fp is not None:
            jsonl_fp.close()

        write_output(out_path, "\n".join(lines), append=False)

        if args.verbose:
            eprint("[DONE] Processed %d files, errors: %d" % (processed, errors))
            eprint("[DONE] Markdown written to: %s" % ("[stdout]" if out_path == "-" else str(Path(out_path).resolve())))
            if jsonl_path is not None:
                eprint("[DONE] JSONL written to: %s" % jsonl_path)

        return

    # Single-file mode
    in_p = Path(args.in_path).resolve()
    if not in_p.exists() or not in_p.is_file():
        eprint("[ERROR] Input file not found: %s" % in_p)
        sys.exit(1)

    kind = classify_ext(in_p)
    if kind == "unknown" and not args.include_unknown:
        eprint("[ERROR] Unknown extension for input file. Use --include-unknown to emit a stub.")
        sys.exit(2)

    try:
        st = in_p.stat()
        mtime = datetime.utcfromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M:%SZ")
        size_b = st.st_size
    except Exception:
        mtime = ""
        size_b = -1

    if args.verbose:
        eprint("[PROC ] %s (%s)" % (in_p.name, kind))

    try:
        content = extract_content(in_p, kind, args.max_pdf_pages, args.clean_text)
    except Exception as e:
        content = "```text\nExtraction error: %s\n```" % e
        if args.verbose:
            eprint("[ERROR] %s: %s" % (in_p.name, e))

    text_block = post_process_block(
        content,
        strip_citations=args.strip_citations,
        remove_pii_flag=args.remove_pii,
        pii_level=args.pii_level,
        max_section_chars=args.max_section_chars,
    )

    section_md = build_section(
        title=in_p.name,
        path_display=in_p.name,
        kind=kind,
        size_b=size_b,
        mtime=mtime,
        text_block=text_block,
        include_divider=False,
        include_metadata=True,
    )

    plain = text_block
    if kind in {"text", "docx", "pdf", "json"}:
        plain = strip_md_fences(plain)

    if args.jsonl_out:
        jsonl_path = Path(args.jsonl_out).resolve()
        jsonl_path.parent.mkdir(parents=True, exist_ok=True)
        mode = "a" if args.append else "w"
        with jsonl_path.open(mode, encoding="utf-8") as jsonl_fp:
            rec = {
                "path": in_p.name,
                "type": kind,
                "sha": sha1_of_path(Path(in_p.name)),
                "modified": mtime,
                "size": size_b,
                "text": plain,
            }
            jsonl_fp.write(json.dumps(rec, ensure_ascii=False) + "\n")
        if args.verbose:
            eprint("[DONE] JSONL written to: %s" % jsonl_path)

    if args.emit == "md":
        write_output(out_path, section_md, append=args.append)
        return
    if args.emit == "text":
        write_output(out_path, plain, append=args.append)
        return

    record = build_json_record(
        path_display=in_p.name,
        kind=kind,
        sha=sha1_of_path(Path(in_p.name)),
        mtime=mtime,
        size_b=size_b,
        markdown=section_md,
        plain=plain,
    )
    write_output(out_path, json.dumps(record, ensure_ascii=False), append=args.append)


if __name__ == "__main__":
    main()
